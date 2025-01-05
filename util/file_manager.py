import os
import shutil
import copy
from markitdown import MarkItDown
from . import data_utils
from docx import Document
import xml.etree.ElementTree as ET
# import PyPDF2
from unstructured.partition.pdf import partition_pdf

class FileManager:
    def __init__(self, root_dir="./Data"):
        self.processed_dir = os.path.join(root_dir, "processed")
        self.raw_dir = os.path.join(root_dir, "raw")
        self.originals_dir = os.path.join(root_dir, "originals")
        self.subdir_dict = {}
        self.raw_files_dict = {}
        self.original_files_dict = {}

    # Reset the processed directory
    def reset_processed_directory(self):
        # Check if the directory exists
        if os.path.exists(self.processed_dir):
            # Delete the directory and all its contents
            shutil.rmtree(self.processed_dir)
        
        # Create a new directory
        os.makedirs(self.processed_dir)

    # Reset the raw directory
    def reset_raw_directory(self):
        # Check if the directory exists
        if os.path.exists(self.raw_dir):
            # Delete the directory and all its contents
            shutil.rmtree(self.raw_dir)
        
        # Create a new directory
        os.makedirs(self.raw_dir)

    # Reset the originals directory
    def reset_originals_directory(self):
        # Check if the directory exists
        if os.path.exists(self.originals_dir):
            # Delete the directory and all its contents
            shutil.rmtree(self.originals_dir)
        
        # Create a new directory
        os.makedirs(self.originals_dir)

        # Create subdirectories for different file types and save the addresses in a dictionary
        subdirs = ["zip", "pdf", "csv", "json", "xlsx", "docx", "png", "jpg", "txt", "other"]
        self.subdir_dict = {}
        for subdir in subdirs:
            subdir_path = os.path.join(self.originals_dir, subdir)
            os.makedirs(subdir_path)
            self.subdir_dict[subdir] = subdir_path

    # Reset all directories
    def reset_all_directories(self):
        self.reset_processed_directory()
        self.reset_raw_directory()
        self.reset_originals_directory()

    # List the files in the raw directory to be processed in a dictionary based on their extensions
    def list_raw_files(self):
        return self.append_raw_files(self.raw_dir)
    
    # Append the files in a folder to the list raw files
    def append_raw_files(self, raw_dir, raw_files_dict=None):
        if raw_files_dict is None:
            raw_files_dict = self.raw_files_dict
        for subdir, dirs, files in os.walk(raw_dir):
            for file in files:
                file_path = os.path.join(subdir, file)
                extension = file.split(".")[-1]
                if extension in raw_files_dict:
                    raw_files_dict[extension].append(file_path)
                else:
                    raw_files_dict[extension] = [file_path]
        return raw_files_dict
    
    # Process the files in the raw directory
    def process_raw_dir(self):
        raw_files_dict = self.list_raw_files()
        # Append hard copy of raw_files_dict to original_files_dict
        self.original_files_dict = copy.deepcopy(raw_files_dict)

        processed_files = []
        
        # Search in raw_files for the extensions of compressed files
        compressed_extensions = ["zip"]
        while compressed_extensions:
            extension = compressed_extensions.pop(0)
            if extension in raw_files_dict:
                compressed_files = raw_files_dict.pop(extension)
                processed_files += self.process_raw_files(compressed_files, extension)

        # Process remaining files in raw_files_dict
        while raw_files_dict:
            extension, files = raw_files_dict.popitem()
            processed_files += self.process_raw_files(files, extension)

        return processed_files
    
    # Process a raw file list based on its extension
    def process_raw_files(self, files, extension):
        switch = {
            "zip": self.process_zip_files,
            "pdf": self.process_pdf_files,
            "csv": self.process_csv_files,
            "json": self.process_json_files,
            "xlsx": self.process_xlsx_files,
            "docx": self.process_docx_files,
            "png": self.process_png_files,
            "jpg": self.process_jpg_files,
            "txt": self.process_txt_files,
            "other": self.process_other_files
        }
        process_function = switch.get(extension, self.process_other_files)
        return process_function(files)

    # Placeholder methods for processing different file types
    def process_zip_files(self, files):

        processed_zip_files = []
        # loop over the file in files     
        while files:
            file = files.pop(0)
            new_raw_files_dict = {}
            #get the path of the file and file name
            file_path = os.path.dirname(file)
            file_name_extension = os.path.basename(file)
            file_name = os.path.splitext(file_name_extension)[0]
            # create a folder with the name of the file with extension
            new_dir = os.path.join(file_path, file_name)
            os.makedirs(new_dir, exist_ok=True)

            try:
                # extract the file
                shutil.unpack_archive(file, new_dir)
                self.append_raw_files(new_dir, raw_files_dict=new_raw_files_dict)
                if "zip" in new_raw_files_dict:
                    files.extend(new_raw_files_dict["zip"])
                    new_raw_files_dict.pop("zip")
                data_utils.append_dictionaries(self.raw_files_dict, new_raw_files_dict)
            except:
                print(f"Error extracting file: {file}")
                continue

            #if the zip file is available in the original_files_dict
            # move the file to the processed directory
            # else delete the zip file
            if file in self.original_files_dict["zip"]:
                shutil.move(file, self.subdir_dict["zip"])
            else:
                os.remove(file)
            
            processed_zip_files.append(file)
        
        return processed_zip_files

    def process_pdf_files(self, files):
        # Implement processing logic for pdf files
        # loop over the file in files
        processed_pdf_files = []
        while files:
            file = files.pop(0)
            main_file_path, main_file_name = os.path.split(file)
            output_dir = os.path.join(self.processed_dir, main_file_name.replace(".","_")+"_images")

            rpe = partition_pdf(
                filename=file,
                strategy="auto",
                extract_image_block_types=["Image", "Table"],
                infer_table_structure=False,
                # chunking_strategy="title",
                max_characters=4000,
                new_after_n_chars=3800,
                combine_text_under_n_chars=2000,
                extract_image_block_output_dir = f"{output_dir}",
                # Add a callback function to update the progress bar
            #   progress_callback=lambda current_page: pbar.update(1) if current_page else None
            )

            new_file_path = os.path.join(self.processed_dir, os.path.splitext(main_file_name)[0] + ".md")

             # Open the markdown file for writing
            with open(new_file_path, 'w', encoding='utf-8') as md_file:
                for el in rpe:
                    if el.category == 'Image' or el.category == 'Figure' or el.category == 'Picture':
                        md_file.write(f"![Image]({el.metadata.image_path})\n")
                    else:
                        md_file.write(f"{el.text}\n")
            processed_pdf_files.append(file)
            # if the pdf file is available in the original_files_dict
            # move the file to the processed directory
            # else delete the pdf file
            if file in self.original_files_dict["pdf"]:
                shutil.move(file, self.subdir_dict["pdf"])
            else:
                os.remove(file)
        return processed_pdf_files

    def process_csv_files(self, files):
        # Implement processing logic for csv files
        pass

    def process_json_files(self, files):
        # Implement processing logic for json files
        pass

    def process_xlsx_files(self, files):
        # Implement processing logic for xlsx files
        pass

    def process_docx_files(self, files):
        namespace = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'o': 'urn:schemas-microsoft-com:office:office',
            'v': 'urn:schemas-microsoft-com:vml',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        processed_docx_files = []
        while files:
            file = files.pop(0)
            # Read docx file and separate content
            doc = Document(file)
            elements = []
            para_start_index = 0
            objects_dict = {}
            obtype_dict = {}
            embedding_objects_types = []  
            main_file_path, main_file_name = os.path.split(file)
           

            '''
            todo: currently, the images and emeding objects are saved in the processed directory
            need to save them in the original directory after processing them into markdown
            # embed_dir = os.path.join(self.subdir_dict["docx"],main_file_name.replace(".","_")+"_embed")
            '''
            embed_dir = os.path.join(self.processed_dir,main_file_name.replace(".","_")+"_embed")

            # delete embeding folder if exists
            if os.path.exists(embed_dir):
                shutil.rmtree(embed_dir)
            os.makedirs(embed_dir, exist_ok=True)
            # save all the embedded objects in the document
            for rel in doc.part.rels.values():
                path, file_name = os.path.split(rel.target_ref)
                path = os.path.split(path)[-1]
                path = os.path.join(embed_dir, path)
                
                os.makedirs(path, exist_ok=True)

                with open(os.path.join(path,file_name), "wb") as f:
                    f.write(rel.target_part.blob)

                objects_dict[rel._rId] = os.path.join(path,file_name)
                obtype_dict[rel._rId] = rel._reltype.split("/")[-1] 

            # Iterate through document elements
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    for index, para in enumerate(doc.paragraphs[para_start_index:]):
                        if para._element == element:
                            para_start_index += index + 1 # Update the start index for the next iteration
                            para_text = para.text
                            if para.style is None:
                                para_style = ""
                            else:
                                para_style = para.style.name.lower()

                            # Check if the paragraph contains an image or text
                            if para_text != "":
                                # Append the paragraph to the elements list
                                elements.append({"type": para_style, "content": para_text})
                            
                            else:
                                xmlstr = str(element.xml)
                                root = ET.fromstring(xmlstr)
                                
                                 # Retrieve r:embed for images
                                blip = root.find('.//a:blip', namespaces=namespace)
                                if blip is not None:
                                    r_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    elements.append({"type": obtype_dict[r_id], "content": objects_dict[r_id]})
                                    # add obtype_dict[r_id] to embedding_objects_types list if not available
                                    if obtype_dict[r_id] not in embedding_objects_types:
                                        embedding_objects_types.append(obtype_dict[r_id])
                                
                                # Retrieve r:id for OLEObjects
                                ole_object = root.find('.//o:OLEObject', namespaces=namespace)
                                if ole_object is not None:
                                    r_id = ole_object.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    elements.append({"type": obtype_dict[r_id], "content": objects_dict[r_id]})
                                    # add obtype_dict[r_id] to embedding_objects_types list if not available
                                    if obtype_dict[r_id] not in embedding_objects_types:
                                        embedding_objects_types.append(obtype_dict[r_id])
                                
                                # Retrieve r:id for v:imagedata
                                imagedata = root.find('.//v:imagedata', namespaces=namespace)
                                if imagedata is not None:
                                    r_id = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    elements.append({"type": obtype_dict[r_id], "content": objects_dict[r_id]})
                                    # add obtype_dict[r_id] to embedding_objects_types list if not available
                                    if obtype_dict[r_id] not in embedding_objects_types:
                                        embedding_objects_types.append(obtype_dict[r_id])


                elif element.tag.endswith('tbl'):
                    # Find the table in doc.tables
                    for table in doc.tables:
                        if table._element == element:
                            table_data = []
                            for row in table.rows:
                                row_data = [cell.text for cell in row.cells]
                                table_data.append(row_data)
                            elements.append({"type": "table", "content": table_data})
                            break
                elif element.tag.endswith('sectPr'):
                    # Handle section properties if needed
                    pass

            #save the content of elements to a markdown file        
            # Construct the new file path in the processed directory
            new_file_path = os.path.join(self.processed_dir, os.path.splitext(main_file_name)[0] + ".md")
            
            # Ensure the directory exists
            os.makedirs(os.path.dirname(new_file_path), exist_ok=True)

            try:
                # Save the elements to the markdown file
                save_elements_to_file(elements, new_file_path, embedding_objects_types=embedding_objects_types)
                processed_docx_files.append(file)
                shutil.move(file, self.subdir_dict["docx"])


            except Exception as e:
                print(f"Error saving file: {new_file_path}")
                print(f"Exception: {e}")
                continue
        
        return processed_docx_files
    
                    
    def process_png_files(self, files):
        # Implement processing logic for png files
        pass

    def process_jpg_files(self, files):
        # Implement processing logic for jpg files
        pass

    def process_txt_files(self, files):
        # Implement processing logic for txt files
        pass

    def process_other_files(self, files):
        # Implement processing logic for other files
        pass

def save_elements_to_file(elements, new_file_path, embedding_objects_types = ["image"]):
    with open(new_file_path, 'w') as f:
        for element in elements:
            element_type = element["type"].lower()
            if ("list" in element_type) or ("bullet" in element_type):
                bullet = '-' if 'bullet' in element_type else '1.'
                f.write(f"{bullet} {element['content']}\n")
            elif element_type == "title":
                f.write("# " + element["content"] + "\n")
            elif element_type.startswith("heading"):
                heading_level = element_type.replace("heading", "")
                f.write("#" * int(heading_level) + " " + element["content"] + "\n")
            elif element_type == "table":
                table_content = element["content"]
                # Write table header
                header = table_content[0]
                f.write("| " + " | ".join(header) + " |\n")
                f.write("|" + " --- |" * len(header) + "\n")
                # Write table rows
                for row in table_content[1:]:
                    f.write("| " + " | ".join(row) + " |\n")
            elif element_type in embedding_objects_types:
                # check if the images folder exists and create it if it does not
                f.write(f"![{element['type']}]({element['content']})\n")
            else:
                # Uncomment if you want to have the custom type/style in the markdown file
                # f.write(f"<!-- Start Custom Type/style: {element["type"]} -->\n")
                f.write(element["content"] + "\n")
                # f.write(f"<!-- End Custom Type/style: {element["type"]} -->\n")

# Example usage
if __name__ == "__main__":
    file_manager = FileManager()
    
    # file_manager.reset_all_directories()
    file_manager.reset_originals_directory()
    # raw_files = file_manager.list_raw_files()
    # print(raw_files)

    processed_files = file_manager.process_raw_dir()